import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
from datetime import date
import io


def _log(event_type, action, detail=""):
    try:
        from module7_audit import log_event
        log_event(event_type, action, detail, module="module5")
    except Exception:
        pass


# ══════════════════════════════════════════════════════
#  WORD REPORT GENERATOR  (unchanged)
# ══════════════════════════════════════════════════════
def generate_docx(bank_name, level, today, summary, gaps):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Cm(2.5)
        sec.top_margin  = sec.bottom_margin = Cm(2)

    NAVY   = RGBColor(0x1e, 0x3a, 0x5f)
    GOLD   = RGBColor(0x92, 0x66, 0x0a)
    BLACK  = RGBColor(0x0f, 0x17, 0x2a)
    GREY   = RGBColor(0x47, 0x55, 0x69)
    GREEN  = RGBColor(0x15, 0x80, 0x3d)
    AMBER  = RGBColor(0xb4, 0x53, 0x09)
    RED    = RGBColor(0xb9, 0x1c, 0x1c)
    PURPLE = RGBColor(0x7c, 0x3a, 0xed)

    pct        = summary["percent"]
    risk       = summary["risk"]
    risk_color = GREEN if pct >= 80 else AMBER if pct >= 60 else RED
    is_lv4     = (level == "Level-IV")

    def heading(text, lvl=1, color=None):
        p = doc.add_heading(text, level=lvl)
        if p.runs:
            p.runs[0].font.color.rgb = color or NAVY
            p.runs[0].font.bold = True
        return p

    def para(text, color=None, bold=False, size=10, align=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size      = Pt(size)
        r.font.bold      = bold
        r.font.color.rgb = color or BLACK
        if align:
            p.alignment = align
        return p

    def shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    t = doc.add_heading("RBI CYBERSECURITY GAP ANALYSIS REPORT", 0)
    if t.runs:
        t.runs[0].font.color.rgb = NAVY
        t.runs[0].font.size = Pt(18)

    para("Primary Urban Cooperative Banks — RBI Graded Approach Framework", color=GREY, size=10)
    if is_lv4:
        para("Comprehensive Framework — All Annexes (Basic + I + II + III + IV) Applicable",
             color=PURPLE, size=9, bold=True)
    doc.add_paragraph()

    heading("Bank Details", 1)
    dt = doc.add_table(rows=2, cols=4)
    dt.style = "Table Grid"
    labels = ["Bank Name", "Compliance Level", "Report Date", "Risk Category"]
    vals   = [bank_name, level, today, risk]
    for i, (lbl, val) in enumerate(zip(labels, vals)):
        row_i, col_i = divmod(i, 2)
        cell_lbl = dt.rows[row_i].cells[col_i * 2]
        cell_val = dt.rows[row_i].cells[col_i * 2 + 1]
        cell_lbl.text = lbl
        cell_val.text = val
        shade_cell(cell_lbl, "EFF6FF")
        cell_lbl.paragraphs[0].runs[0].font.bold      = True
        cell_lbl.paragraphs[0].runs[0].font.color.rgb = NAVY
        cell_lbl.paragraphs[0].runs[0].font.size      = Pt(9)
        cell_val.paragraphs[0].runs[0].font.size       = Pt(9)
        if is_lv4 and i == 1:
            cell_val.paragraphs[0].runs[0].font.color.rgb = PURPLE
    doc.add_paragraph()

    heading("Compliance Score", 1)
    sc = doc.add_table(rows=1, cols=3)
    sc.style = "Table Grid"
    sc.rows[0].cells[0].text = f"Score: {pct}%"
    sc.rows[0].cells[1].text = f"Risk: {risk}"
    if pct >= 80:   verdict_txt = "Strong compliance posture."
    elif pct >= 60: verdict_txt = "Moderate compliance — gaps require attention."
    else:           verdict_txt = "Critical deficiencies — immediate action required."
    sc.rows[0].cells[2].text = verdict_txt
    shade_clr = "F0FDF4" if pct >= 80 else "FFFBEB" if pct >= 60 else "FEF2F2"
    for i, cell in enumerate(sc.rows[0].cells):
        shade_cell(cell, shade_clr)
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs \
              else cell.paragraphs[0].add_run(cell.text)
        run.font.bold      = (i < 2)
        run.font.color.rgb = risk_color
        run.font.size      = Pt(10 if i < 2 else 9)
    doc.add_paragraph()

    heading("Compliance Metrics", 1)
    mt = doc.add_table(rows=1, cols=2)
    mt.style = "Table Grid"
    for cell, txt in zip(mt.rows[0].cells, ["Metric", "Value"]):
        cell.text = txt
        shade_cell(cell, "1E3A5F")
        r = cell.paragraphs[0].runs[0]
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.bold = True
        r.font.size = Pt(9)

    metrics = [
        ("Total Controls Assessed", str(summary["total"])),
        ("Compliant Controls",       str(summary["compliant"])),
        ("Partially Compliant",      str(summary["partial"])),
        ("Non-Compliant Controls",   str(summary["non"])),
        ("Not Applicable",           str(summary["na"])),
        ("Not Assessed (Unfilled)",  str(summary.get("not_assessed", 0))),
        ("Compliance Score",         f"{pct}%"),
        ("Risk Category",            risk),
    ]


    for i, (lbl, val) in enumerate(metrics):
        row = mt.add_row()
        row.cells[0].text = lbl
        row.cells[1].text = val
        fill = "F8FAFC" if i % 2 == 0 else "FFFFFF"
        shade_cell(row.cells[0], fill)
        shade_cell(row.cells[1], fill)
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    if is_lv4:
        heading("Annex IV — Governance Controls Summary", 1, color=PURPLE)
        gov_items = [
            ("C-SOC (Cyber Security Operation Centre)", "Must be set up with 24x7 surveillance capability and SIEM integration (Sections 1.1 & 1.2)."),
            ("Cyber Drill Participation", "UCB must participate in CERT-IN / IDRBT cyber drills (Section 2.1)."),
            ("Incident Response Framework", "Incident response capability across all interconnected systems; SOC, IR and Digital Forensics policy required (Sections 3.1 & 3.2)."),
            ("Forensics & Metrics", "KPIs/KRIs for security operations; network forensics and DDoS mitigation arrangement on standby (Sections 4.1 & 4.2)."),
            ("IT Strategy & Policy", "Board-approved IT strategy covering hardware/software architecture, outsourcing, IT org structure, and training (Section 5.1)."),
            ("Cyber Security Team / Function", "Dedicated cyber security function with adequate staffing and tools (Section 6.1)."),
            ("IT Strategy Committee", "Board-level committee with minimum 2 directors; at least 2 technically competent members (Section 6.2)."),
            ("IT Steering Committee", "Executive-level committee for IT strategy implementation, project prioritisation and compliance (Section 6.3)."),
            ("CISO Appointment", "Sufficiently senior CISO reporting directly to risk management head or CEO; must not have business targets (Section 6.4)."),
            ("Information Security Committee", "Quarterly meetings; CISO as member secretary; must report to Board on cyber security activities (Section 6.5)."),
            ("Audit Committee of Board (ACB)", "ACB responsible for IS Audit performance evaluation and VA-PT compliance monitoring (Section 6.6)."),
        ]
        gt = doc.add_table(rows=1, cols=2)
        gt.style = "Table Grid"
        for cell, txt in zip(gt.rows[0].cells, ["Governance Control", "Requirement"]):
            cell.text = txt
            shade_cell(cell, "3B1F6E")
            r = cell.paragraphs[0].runs[0]
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True
            r.font.size = Pt(9)
        for ri, (ctrl, req) in enumerate(gov_items):
            row = gt.add_row()
            row.cells[0].text = ctrl
            row.cells[1].text = req
            fill = "F5F0FF" if ri % 2 == 0 else "FFFFFF"
            shade_cell(row.cells[0], fill)
            shade_cell(row.cells[1], fill)
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
        doc.add_paragraph()

    if gaps is not None and not gaps.empty:
        heading("Key Compliance Gaps", 1)
        sec_col = next((c for c in gaps.columns if "section"     in c.lower()), None)
        req_col = next((c for c in gaps.columns if "requirement" in c.lower()), None)
        sc_col  = "_status"       if "_status"       in gaps.columns else None
        sl_col  = "_source_label" if "_source_label" in gaps.columns else \
                  "_source"       if "_source"       in gaps.columns else None
        pc_col  = "Priority"      if "Priority"      in gaps.columns else None
        show_c  = [c for c in [sec_col, req_col, sc_col, sl_col, pc_col]
                   if c and c in gaps.columns]
        if show_c:
            gap_show = gaps[show_c].rename(columns={
                "_status": "Status", "_source_label": "Annex", "_source": "Annex"
            }).head(30)
            gt2 = doc.add_table(rows=1, cols=len(gap_show.columns))
            gt2.style = "Table Grid"
            for i, col in enumerate(gap_show.columns):
                cell = gt2.rows[0].cells[i]
                cell.text = col
                shade_cell(cell, "1E3A5F")
                r = cell.paragraphs[0].runs[0]
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(8)
            for ri, (_, row) in enumerate(gap_show.iterrows()):
                tr   = gt2.add_row()
                fill = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
                for i, val in enumerate(row):
                    tr.cells[i].text = str(val) if val is not None else ""
                    shade_cell(tr.cells[i], fill)
                    if tr.cells[i].paragraphs[0].runs:
                        tr.cells[i].paragraphs[0].runs[0].font.size = Pt(8)
        doc.add_paragraph()

    heading("Recommendations", 1)
    recs = []
    if summary["non"]     > 0:
        recs.append(f"Immediately remediate {summary['non']} non-compliant controls identified in the Gap Analysis.")
    if summary["partial"] > 0:
        recs.append(f"Develop a 90-day action plan for {summary['partial']} partially compliant controls.")
    if pct < 80:
        recs.append("Conduct a cybersecurity awareness and training program for all relevant staff.")
        recs.append("Assign a dedicated Cybersecurity Officer to oversee remediation tracking.")
        recs.append("Perform quarterly internal compliance re-assessments and maintain updated evidence artefacts.")
    if is_lv4:
        recs.append("Establish the Cyber Security Operation Centre (C-SOC) with 24x7 SIEM-integrated surveillance capability at the earliest.")
        recs.append("Appoint a dedicated CISO who reports directly to the risk management head or CEO, with no business targets assigned.")
        recs.append("Constitute the IT Strategy Committee (Board level) and IT Steering Committee with technically competent members as mandated under Annex IV.")
        recs.append("Form the Information Security Committee with quarterly review meetings; CISO must serve as member secretary.")
        recs.append("Ensure the Audit Committee of Board (ACB) is equipped to evaluate IS Audit findings and monitor VA-PT compliance.")
        recs.append("Register and participate in cyber drills conducted by CERT-IN and IDRBT as required under Annex IV Section 2.")
        recs.append("Implement a documented Incident Response and Digital Forensics framework covering all interconnected systems and vendor networks.")
        recs.append("Develop a Board-approved IT Strategy covering hardware/software architecture, outsourcing strategy, and IT Governance structure.")
    recs.append("Submit compliance status reports to RBI as per prescribed timelines and guidelines.")
    for i, rec in enumerate(recs, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(rec)
        r.font.size      = Pt(10)
        r.font.color.rgb = BLACK

    doc.add_paragraph()
    para(
        f"This report was auto-generated on {today} using the RBI CSF Compliance Assessment Tool. "
        "For internal use only. Verify all data against the original filled checklists before submission to RBI.",
        color=GREY, size=8,
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════
#  SHARED HTML HELPERS
# ══════════════════════════════════════════════════════
def _is_light():
    return st.session_state.get("theme", "dark") == "light"


def _theme(light):
    if light:
        return dict(
            bg="#ffffff", bg2="#f8fafc", border="#e2e8f0",
            text="#0f172a", text2="#334155", dim="#94a3b8",
            gold="#8b5c0a", font="IBM Plex Sans, system-ui, sans-serif",
        )
    return dict(
        bg="#08121f", bg2="#0d1a30", border="#162540",
        text="#dde4f0", text2="#8a9ab8", dim="#4a5e7a",
        gold="#c9a84c", font="IBM Plex Sans, system-ui, sans-serif",
    )


def _wrap(css_body, html_body):
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<style>body{margin:0;padding:0;background:transparent;}"
        + css_body +
        "</style></head><body>"
        + html_body +
        "</body></html>"
    )


# ── 1. Page header ────────────────────────────────────
def _page_header_html(light):
    t = _theme(light)
    css = (
        f".hdr{{background:linear-gradient(135deg,{t['bg']} 0%,{t['bg2']} 100%);"
        f"border:1px solid {t['border']};border-radius:16px;padding:28px 32px;"
        f"border-left:4px solid {t['gold']};margin-bottom:4px;}}"
        f"h1{{font-family:Rajdhani,sans-serif;font-size:26px;font-weight:700;"
        f"color:{t['text']};margin:0 0 6px 0;letter-spacing:0.5px;}}"
        f"p{{font-size:13px;color:{t['dim']};margin:0;}}"
    )
    html = (
        '<div class="hdr">'
        "<h1>&#128202; MODULE 5 &mdash; GAP ANALYSIS REPORT</h1>"
        "<p>Review your compliance posture and export the gap analysis report in Word format</p>"
        "</div>"
    )
    return _wrap(css, html)


# ── 2. Report header card ─────────────────────────────
def _report_header_html(bank_name, bank_level, today, is_lv4, light):
    t     = _theme(light)
    acc   = "#a855f7" if is_lv4 else t["gold"]
    badge_bg  = {"Level-I":"#1e3a5f","Level-II":"#78350f","Level-III":"#7f1d1d","Level-IV":"#4c1d95"}.get(bank_level,"#1e3a5f")
    badge_txt = "#93c5fd"
    lv4_line  = (
        f'<div style="font-size:11px;color:{acc};margin-top:5px;font-weight:600;">'
        '&#128217; Full Framework &mdash; Basic + Annex I + II + III + IV</div>'
    ) if is_lv4 else ""
    css = (
        f"*{{font-family:{t['font']};box-sizing:border-box;}}"
        f".card{{background:{t['bg']};border:1px solid {acc};border-radius:14px;"
        f"padding:22px 26px;display:flex;justify-content:space-between;"
        f"align-items:flex-start;flex-wrap:wrap;gap:14px;}}"
        f".name{{font-family:Rajdhani,sans-serif;font-size:22px;font-weight:700;color:{t['text']};margin:0;}}"
        f".sub{{font-size:11px;color:{t['dim']};margin-top:3px;letter-spacing:0.8px;text-transform:uppercase;}}"
        f".right{{text-align:right;}}"
        f".lbl{{font-size:9.5px;font-weight:700;letter-spacing:0.9px;text-transform:uppercase;color:{t['dim']};margin-bottom:4px;}}"
        f".date{{font-family:'IBM Plex Mono',monospace;font-size:14px;color:{t['gold']};font-weight:600;}}"
        f".badge{{display:inline-block;padding:3px 12px;border-radius:20px;font-size:11px;"
        f"font-weight:700;letter-spacing:0.5px;background:{badge_bg};color:{badge_txt};margin-top:6px;}}"
    )
    html = (
        '<div class="card">'
        '<div>'
        f'<div class="name">&#127970; {bank_name}</div>'
        '<div class="sub">RBI Cybersecurity &nbsp;&middot;&nbsp; Gap Analysis Report</div>'
        + lv4_line +
        '</div>'
        '<div class="right">'
        '<div class="lbl">Report Date</div>'
        f'<div class="date">{today}</div>'
        f'<div><span class="badge">{bank_level}</span></div>'
        '</div>'
        '</div>'
    )
    return _wrap(css, html)


# ── 3. KPI stat strip ────────────────────────────────
def _stat_strip_html(items, light):
    """items = list of (label, value, color_or_None)"""
    t = _theme(light)
    css = (
        f"*{{font-family:{t['font']};box-sizing:border-box;}}"
        f".strip{{background:{t['bg']};border:1px solid {t['border']};border-radius:14px;"
        f"padding:14px 8px;display:flex;align-items:center;flex-wrap:wrap;}}"
        f".item{{flex:1;min-width:80px;padding:6px 14px;border-right:1px solid {t['border']};text-align:center;}}"
        f".item:last-child{{border-right:none;}}"
        f".lbl{{font-size:9.5px;font-weight:700;letter-spacing:0.9px;text-transform:uppercase;color:{t['dim']};margin-bottom:4px;}}"
        f".val{{font-family:Rajdhani,sans-serif;font-size:26px;font-weight:700;line-height:1.1;color:{t['gold']};}}"
    )
    cells = ""
    for label, value, color in items:
        c = color if color else t["gold"]
        cells += (
            '<div class="item">'
            f'<div class="lbl">{label}</div>'
            f'<div class="val" style="color:{c};">{value}</div>'
            '</div>'
        )
    return _wrap(css, '<div class="strip">' + cells + "</div>")


# ── 4. Score banner ──────────────────────────────────
def _score_banner_html(pct, risk, verdict_html, risk_color, risk_badge_label, light):
    t = _theme(light)
    badge_bg = {"LOW RISK":"#14532d","MODERATE RISK":"#78350f","HIGH RISK":"#7f1d1d"}.get(risk_badge_label, "#1e3a5f")
    badge_tc = {"LOW RISK":"#86efac","MODERATE RISK":"#fde68a","HIGH RISK":"#fca5a5"}.get(risk_badge_label, "#93c5fd")
    css = (
        f"*{{font-family:{t['font']};box-sizing:border-box;}}"
        f".card{{background:{t['bg']};border:1px solid {t['border']};border-radius:14px;"
        f"padding:20px 24px;display:flex;align-items:center;gap:24px;flex-wrap:wrap;}}"
        f".divider{{width:1px;height:60px;background:{t['border']};flex-shrink:0;}}"
        f".lbl{{font-size:9.5px;font-weight:700;letter-spacing:0.9px;text-transform:uppercase;color:{t['dim']};margin-bottom:6px;}}"
        f".score{{font-family:Rajdhani,sans-serif;font-size:50px;font-weight:700;color:{risk_color};line-height:1;}}"
        f".badge{{display:inline-block;padding:4px 14px;border-radius:20px;font-size:12px;"
        f"font-weight:700;letter-spacing:0.5px;background:{badge_bg};color:{badge_tc};}}"
        f".verdict{{font-size:13px;color:{t['text2']};line-height:1.65;flex:1;min-width:200px;}}"
    )
    html = (
        '<div class="card">'
        '<div><div class="lbl">Compliance Score</div>'
        f'<div class="score">{pct}%</div></div>'
        '<div class="divider"></div>'
        '<div><div class="lbl">Risk Category</div>'
        f'<div><span class="badge">{risk_badge_label}</span></div></div>'
        '<div class="divider"></div>'
        f'<div class="verdict"><div class="lbl">Assessment Verdict</div>{verdict_html}</div>'
        '</div>'
    )
    return _wrap(css, html)


# ── 5. Governance checklist card (single item) ────────
def _gov_card_html(title, desc, light):
    t = _theme(light)
    css = (
        f"*{{font-family:{t['font']};box-sizing:border-box;}}"
        f".card{{background:{t['bg']};border:1px solid {t['border']};"
        f"border-left:3px solid #a855f7;border-radius:10px;padding:11px 14px;margin-bottom:2px;}}"
        f".title{{font-size:12px;font-weight:700;color:#a855f7;margin-bottom:3px;}}"
        f".desc{{font-size:11.5px;color:{t['text2']};line-height:1.5;}}"
    )
    html = (
        '<div class="card">'
        f'<div class="title">&#128217; {title}</div>'
        f'<div class="desc">{desc}</div>'
        '</div>'
    )
    return _wrap(css, html)


# ── 6. Recommendations list ──────────────────────────
def _recs_html(recs, light):
    t = _theme(light)
    items_html = "".join(
        f'<li style="margin-bottom:9px;line-height:1.65;">{r}</li>'
        for r in recs
    )
    css = (
        f"*{{font-family:{t['font']};box-sizing:border-box;}}"
        f".card{{background:{t['bg']};border:1px solid {t['border']};border-radius:14px;padding:18px 22px;}}"
        f"ol{{margin:0;padding-left:20px;color:{t['text2']};font-size:13px;}}"
    )
    html = f'<div class="card"><ol>{items_html}</ol></div>'
    return _wrap(css, html)


# ── 7. Export card ───────────────────────────────────
def _export_card_html(is_lv4, light):
    t   = _theme(light)
    acc = "#a855f7" if is_lv4 else t["gold"]
    lv4_note = (
        f' &nbsp;&middot;&nbsp; <strong style="color:#a855f7;">Annex IV governance section included</strong>'
    ) if is_lv4 else ""
    icon = "&#128217;" if is_lv4 else "&#128216;"
    css = (
        f"*{{font-family:{t['font']};box-sizing:border-box;}}"
        f".card{{background:{t['bg']};border:1px solid {acc};border-radius:14px;"
        f"padding:16px 20px;display:flex;align-items:center;gap:16px;}}"
        f".icon{{font-size:30px;flex-shrink:0;}}"
        f".title{{font-family:Rajdhani,sans-serif;font-size:16px;font-weight:600;color:{t['text']};margin-bottom:3px;}}"
        f".sub{{font-size:12px;color:{t['text2']};}}"
    )
    html = (
        '<div class="card">'
        f'<div class="icon">{icon}</div>'
        '<div>'
        '<div class="title">Word Report</div>'
        f'<div class="sub">Fully editable .docx &nbsp;&middot;&nbsp; Coloured tables &nbsp;&middot;&nbsp; Compliance summary{lv4_note}</div>'
        '</div>'
        '</div>'
    )
    return _wrap(css, html)


# ══════════════════════════════════════════════════════
#  MODULE 5 UI
# ══════════════════════════════════════════════════════
def show_module5():

    light = _is_light()

    # Page header
    st.markdown("""
    <div class="rbi-page-header">
        <h1>📊 MODULE 5 — GAP ANALYSIS REPORT</h1>
        <p>Review your compliance posture and export the gap analysis report in Word format</p>
    </div>
    """, unsafe_allow_html=True)

    _log("view", "Gap Analysis Report viewed")

    summary    = st.session_state.get("compliance_summary")
    gaps       = st.session_state.get("gap_dataframe")
    bank_name  = st.session_state.get("bank_name",  "Not Specified")
    bank_level = st.session_state.get("bank_level", "Not Set")
    is_lv4     = (bank_level == "Level-IV")

    if not summary:
        st.warning("Please complete Module 3 (Compliance Dashboard) first to generate the report.")
        return

    today = date.today().strftime("%d %B %Y")
    pct   = summary["percent"]
    risk  = summary["risk"]

    risk_color = {"Low Risk": "#15803d", "Moderate Risk": "#b45309", "High Risk": "#b91c1c"}[risk]
    risk_badge = risk.upper()

    if pct >= 80:
        verdict_html = f"The bank demonstrates a <strong style='color:#15803d;'>strong compliance posture</strong> with the RBI Cybersecurity Framework."
    elif pct >= 60:
        verdict_html = f"The bank shows <strong style='color:#b45309;'>moderate compliance</strong> but has significant gaps requiring attention."
    else:
        verdict_html = f"The bank has <strong style='color:#b91c1c;'>critical compliance deficiencies</strong> that require immediate remediation."

    # ── Report header card ─────────────────────────────
    components.html(_report_header_html(bank_name, bank_level, today, is_lv4, light),
                    height=110, scrolling=False)

    # ── KPI stat strip ─────────────────────────────────
    st.markdown('<div class="rbi-section-title">Compliance Summary</div>', unsafe_allow_html=True)
    strip_items = [
        ("Total Controls",  summary["total"],                    None),
        ("Compliant",       summary["compliant"],                "#15803d"),
        ("Partial",         summary["partial"],                  "#b45309"),
        ("Non-Compliant",   summary["non"],                      "#b91c1c"),
        ("N/A",             summary["na"],                       None),
        ("Not Assessed",    summary.get("not_assessed", 0),     "#475569"),
    ]
    components.html(_stat_strip_html(strip_items, light), height=95, scrolling=False)

    # ── Score banner ───────────────────────────────────
    components.html(
        _score_banner_html(pct, risk, verdict_html, risk_color, risk_badge, light),
        height=112, scrolling=False,
    )
    st.progress(pct / 100)

    # ── Annex IV Governance Checklist ──────────────────
    if is_lv4:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="rbi-section-title">Annex IV — Governance Checklist</div>',
                    unsafe_allow_html=True)
        gov_checks = [
            ("C-SOC Setup",                "Cyber Security Operation Centre with 24x7 SIEM surveillance"),
            ("Cyber Drill Participation",   "CERT-IN / IDRBT cyber drill participation"),
            ("Incident Response Framework", "IR & Digital Forensics policy across all interconnected systems"),
            ("Forensic Metrics",            "KPIs/KRIs for security ops; DDoS mitigation on standby"),
            ("Board-approved IT Strategy",  "Covers architecture, outsourcing, org structure, training plan"),
            ("CISO Appointed",              "Reports to risk head/CEO; no business targets; adequate budget"),
            ("IT Strategy Committee",       "Board level; 2+ technically competent directors"),
            ("IT Steering Committee",       "Executive level; implementation focus; quarterly reporting"),
            ("IS Committee",                "Quarterly meetings; CISO as member secretary"),
            ("ACB IS Audit Role",           "IS audit findings reviewed; VA-PT compliance monitored"),
        ]
        c1, c2 = st.columns(2)
        for i, (title, desc) in enumerate(gov_checks):
            col = c1 if i % 2 == 0 else c2
            with col:
                components.html(_gov_card_html(title, desc, light), height=76, scrolling=False)

    # ── Gaps Preview ───────────────────────────────────
    if gaps is not None and not gaps.empty:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="rbi-section-title">Key Compliance Gaps</div>',
                    unsafe_allow_html=True)

        # Assign priority from _status if not already set by module4
        if "Priority" not in gaps.columns or gaps["Priority"].isna().all():
            def _assign_p(row):
                s = str(row.get("_status", "")).lower()
                if "not compliant" in s:
                    return "High"
                return "Medium"
            gaps = gaps.copy()
            gaps["Priority"] = gaps.apply(_assign_p, axis=1)
            st.session_state["gap_dataframe"] = gaps

        # Normalise emoji-prefixed values written by older sessions
        gaps["Priority"] = gaps["Priority"].astype(str).str.replace(
            r"^[^\w]*", "", regex=True
        ).str.strip().str.replace("Critical", "Critical").str.replace(
            "High", "High"
        )

        high     = int((gaps["Priority"] == "High").sum())
        critical = int((gaps["Priority"] == "Critical").sum())
        med      = int((gaps["Priority"] == "Medium").sum())

        gap_strip = [("Total Gaps", len(gaps), "#b91c1c")]
        if is_lv4 and critical > 0:
            gap_strip.append(("Critical (Gov.)", critical, "#a855f7"))
        gap_strip += [
            ("High Priority",   high, "#ef4444"),
            ("Medium Priority", med,  "#f59e0b"),
        ]
        components.html(_stat_strip_html(gap_strip, light), height=95, scrolling=False)

        sec_col = next((c for c in gaps.columns if "section"     in c.lower()), None)
        req_col = next((c for c in gaps.columns if "requirement" in c.lower()), None)
        sl_col  = "_source_label" if "_source_label" in gaps.columns else \
                  "_source"       if "_source"       in gaps.columns else None
        sc      = "_status"  if "_status"  in gaps.columns else None
        pc      = "Priority" if "Priority" in gaps.columns else None
        show_c  = [c for c in [sec_col, req_col, sc, sl_col, pc] if c and c in gaps.columns]
        if show_c:
            st.dataframe(
                gaps[show_c].rename(columns={
                    "_status": "Status", "_source_label": "Annex", "_source": "Annex"
                }).head(20),
                use_container_width=True,
                height=250,
            )

    # ── Recommendations ────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="rbi-section-title">Recommendations</div>', unsafe_allow_html=True)

    recs = []
    if summary["non"]     > 0:
        recs.append(f"Immediately remediate the <strong>{summary['non']} non-compliant controls</strong> identified in the Gap Analysis.")
    if summary["partial"] > 0:
        recs.append(f"Develop a 90-day action plan for the <strong>{summary['partial']} partially compliant controls</strong>.")
    if pct < 80:
        recs.append("Conduct a cybersecurity awareness and training program for all relevant staff.")
        recs.append("Assign a dedicated Cybersecurity Officer to oversee remediation tracking.")
    if is_lv4:
        recs.append("<strong>[Annex IV]</strong> Establish the <strong>C-SOC</strong> with 24x7 SIEM-integrated surveillance at the earliest.")
        recs.append("<strong>[Annex IV]</strong> Appoint a <strong>dedicated CISO</strong> reporting to the risk head/CEO with no business targets.")
        recs.append("<strong>[Annex IV]</strong> Constitute the <strong>IT Strategy Committee</strong> and <strong>IT Steering Committee</strong> at Board/executive level.")
        recs.append("<strong>[Annex IV]</strong> Form the <strong>Information Security Committee</strong> (quarterly reviews; CISO as secretary) and ensure <strong>ACB</strong> reviews IS audit findings.")
        recs.append("<strong>[Annex IV]</strong> Participate in <strong>CERT-IN / IDRBT Cyber Drills</strong> and document participation records.")
        recs.append("<strong>[Annex IV]</strong> Implement a documented <strong>Incident Response &amp; Digital Forensics</strong> framework covering all interconnected systems.")
        recs.append("<strong>[Annex IV]</strong> Obtain <strong>Board approval for IT Strategy</strong> covering architecture, outsourcing, org structure, and training.")
    recs.append("Perform quarterly internal compliance re-assessments and maintain updated evidence artefacts.")
    recs.append("Submit compliance status reports to RBI as per prescribed timelines and guidelines.")

    rec_height = min(60 + len(recs) * 42, 520)
    components.html(_recs_html(recs, light), height=rec_height, scrolling=False)

    # ── Export ─────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="rbi-section-title">Export Report</div>', unsafe_allow_html=True)
    components.html(_export_card_html(is_lv4, light), height=84, scrolling=False)

    safe_name = bank_name.replace(" ", "_")
    try:
        docx_buf = generate_docx(bank_name, bank_level, today, summary, gaps)
        if st.download_button(
            "Download  Word Report (.docx)",
            docx_buf,
            file_name=f"Gap_Analysis_Report_{safe_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        ):
            _log("export", "Word Report downloaded",
                 detail=f"{bank_name} / {bank_level} / Score: {pct}%")
    except Exception as e:
        st.error(f"Word generation error: {e}")